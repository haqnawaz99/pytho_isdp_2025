"""
Convert CHAPTER_Variables.md to DOCX with embedded images.
Run from the 2_Variables folder: python md_to_docx.py
Requires: pip install python-docx
"""
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx: pip install python-docx")
    raise

CHAPTER_DIR = Path(__file__).resolve().parent
MD_FILE = CHAPTER_DIR / "CHAPTER_Variables.md"
OUT_DOCX = CHAPTER_DIR / "CHAPTER_Variables.docx"
DIAGRAMS_DIR = CHAPTER_DIR / "diagrams"


def parse_md(content):
    """Split content into blocks: heading, paragraph, code, image."""
    blocks = []
    lines = content.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
            continue
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            blocks.append(("code", "\n".join(code_lines)))
            continue
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)\s*$', line)
        if img_match:
            blocks.append(("image", (img_match.group(1), img_match.group(2))))
            i += 1
            continue
        if line.strip() == "---" or line.strip() == "":
            i += 1
            continue
        para_lines = []
        while i < len(lines):
            l = lines[i]
            if l.startswith("#") or l.strip().startswith("```") or re.match(r'^!\[.*?\]\(.*?\)', l.strip()) or l.strip() == "---":
                break
            if l.strip() == "" and para_lines:
                break
            para_lines.append(l)
            i += 1
        text = "\n".join(para_lines).strip()
        if text:
            blocks.append(("para", text))
    return blocks


def add_paragraph_plain(doc, text):
    """Add paragraph with **bold** and `code` styling."""
    p = doc.add_paragraph()
    pos = 0
    while pos < len(text):
        bold_match = re.search(r"\*\*([^*]+)\*\*", text[pos:])
        code_match = re.search(r"`([^`]+)`", text[pos:])
        bstart = bold_match.start() + pos if bold_match else len(text)
        cstart = code_match.start() + pos if code_match else len(text)
        next_special = min(bstart, cstart)
        if next_special > pos:
            p.add_run(text[pos:next_special])
            pos = next_special
        elif bstart <= cstart and bold_match:
            run = p.add_run(bold_match.group(1))
            run.bold = True
            pos = bstart + len(bold_match.group(0))
        elif code_match:
            run = p.add_run(code_match.group(1))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            pos = cstart + len(code_match.group(0))
        else:
            break
    return p


def main():
    if not MD_FILE.exists():
        print(f"Not found: {MD_FILE}")
        return
    content = MD_FILE.read_text(encoding="utf-8")
    blocks = parse_md(content)
    doc = Document()
    doc.add_heading("Chapter: Variables in Python", 0)
    for kind, data in blocks:
        if kind == "h1":
            doc.add_heading(data, level=1)
        elif kind == "h2":
            doc.add_heading(data, level=2)
        elif kind == "h3":
            doc.add_heading(data, level=3)
        elif kind == "para":
            add_paragraph_plain(doc, data)
        elif kind == "code":
            p = doc.add_paragraph()
            p.style = "Normal"
            run = p.add_run(data)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif kind == "image":
            alt, path = data
            img_path = (CHAPTER_DIR / path).resolve()
            if img_path.exists():
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(4.5))
                doc.add_paragraph()
            else:
                doc.add_paragraph(f"[Image: {alt} — file not found: {path}]")
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Saved: {OUT_DOCX}")


if __name__ == "__main__":
    main()
